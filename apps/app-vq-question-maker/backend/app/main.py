from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal
import json
import re
import uuid

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from PIL import Image
import numpy as np
import fitz
import pytesseract
from docx import Document
from docx.shared import Mm
from openpyxl import Workbook

DATA_ROOT = Path("/data/jobs")
DEFAULT_THUMB_SIZE = (240, 320)

app = FastAPI(title="Lesson Prep Agent")


@dataclass
class PageImageInfo:
    path: Path
    width: int
    height: int


class Rect(BaseModel):
    x: float
    y: float
    w: float
    h: float


class SuggestRegionsRequest(BaseModel):
    page_indexes: list[int] = Field(default_factory=list)


class OCRPageRequest(BaseModel):
    page_index: int
    main_rect: Rect
    vocab_rect: Rect


class OCRRequest(BaseModel):
    pages: list[OCRPageRequest]
    ocr_engine: Literal["tesseract", "gcv", "docai"] = "tesseract"
    language_hint: str | None = None


class ExportWordRequest(BaseModel):
    title: str | None = None
    pages: list[dict[str, Any]]
    layout: dict[str, Any] | None = None
    template_id: str | None = None


class ExportExcelRequest(BaseModel):
    vocab_rows_all: list[dict[str, Any]]
    format: Literal["xlsx"] = "xlsx"


class ExportSlidesRequest(BaseModel):
    sentences_all: list[str]
    style: dict[str, Any]
    google: dict[str, Any] | None = None


def ensure_job_dir(job_id: str) -> Path:
    job_dir = DATA_ROOT / job_id
    (job_dir / "pages").mkdir(parents=True, exist_ok=True)
    (job_dir / "thumbs").mkdir(parents=True, exist_ok=True)
    (job_dir / "exports").mkdir(parents=True, exist_ok=True)
    return job_dir


def render_pdf_pages(pdf_path: Path, job_dir: Path) -> list[PageImageInfo]:
    doc = fitz.open(pdf_path)
    pages_info: list[PageImageInfo] = []
    for index in range(doc.page_count):
        page = doc.load_page(index)
        matrix = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=matrix)
        page_path = job_dir / "pages" / f"page_{index:03d}.png"
        pix.save(page_path.as_posix())
        with Image.open(page_path) as img:
            pages_info.append(PageImageInfo(page_path, img.width, img.height))
            thumb = img.copy()
            thumb.thumbnail(DEFAULT_THUMB_SIZE)
            thumb_path = job_dir / "thumbs" / f"thumb_{index:03d}.png"
            thumb.save(thumb_path)
    return pages_info


def density_map(image: Image.Image, grid_size: int = 100) -> np.ndarray:
    gray = image.convert("L")
    small = gray.resize((grid_size, grid_size))
    arr = np.array(small)
    return 1.0 - (arr / 255.0)


def largest_component(mask: np.ndarray) -> tuple[int, int, int, int] | None:
    visited = np.zeros_like(mask, dtype=bool)
    best_bbox = None
    best_area = 0
    height, width = mask.shape

    for y in range(height):
        for x in range(width):
            if not mask[y, x] or visited[y, x]:
                continue
            stack = [(y, x)]
            visited[y, x] = True
            min_y, max_y, min_x, max_x = y, y, x, x
            area = 0
            while stack:
                cy, cx = stack.pop()
                area += 1
                min_y = min(min_y, cy)
                max_y = max(max_y, cy)
                min_x = min(min_x, cx)
                max_x = max(max_x, cx)
                for ny, nx in ((cy - 1, cx), (cy + 1, cx), (cy, cx - 1), (cy, cx + 1)):
                    if 0 <= ny < height and 0 <= nx < width and mask[ny, nx] and not visited[ny, nx]:
                        visited[ny, nx] = True
                        stack.append((ny, nx))
            if area > best_area:
                best_area = area
                best_bbox = (min_x, min_y, max_x, max_y)
    return best_bbox


def suggest_regions_for_image(image: Image.Image) -> dict[str, Rect]:
    grid = density_map(image, grid_size=100)
    threshold = float(np.quantile(grid, 0.85))
    mask = grid > threshold

    height, width = grid.shape
    center_mask = mask.copy()
    center_mask[:, : int(width * 0.1)] = False
    center_mask[:, int(width * 0.9) :] = False
    center_mask[: int(height * 0.05), :] = False
    center_mask[int(height * 0.9) :, :] = False

    main_bbox = largest_component(center_mask)

    right_mask = mask.copy()
    right_mask[:, : int(width * 0.7)] = False
    right_bbox = largest_component(right_mask)

    bottom_mask = mask.copy()
    bottom_mask[: int(height * 0.7), :] = False
    bottom_bbox = largest_component(bottom_mask)

    img_w, img_h = image.size

    def scale_bbox(bbox: tuple[int, int, int, int] | None) -> Rect | None:
        if bbox is None:
            return None
        min_x, min_y, max_x, max_y = bbox
        return Rect(
            x=(min_x / width) * img_w,
            y=(min_y / height) * img_h,
            w=((max_x - min_x + 1) / width) * img_w,
            h=((max_y - min_y + 1) / height) * img_h,
        )

    main_rect = scale_bbox(main_bbox)
    vocab_rect = scale_bbox(right_bbox or bottom_bbox)

    if main_rect is None:
        main_rect = Rect(x=0.1 * img_w, y=0.1 * img_h, w=0.65 * img_w, h=0.7 * img_h)
    if vocab_rect is None:
        vocab_rect = Rect(x=0.75 * img_w, y=0.1 * img_h, w=0.22 * img_w, h=0.8 * img_h)

    padding_x = img_w * 0.01
    padding_y = img_h * 0.01

    def pad(rect: Rect) -> Rect:
        x = max(rect.x - padding_x, 0)
        y = max(rect.y - padding_y, 0)
        w = min(rect.w + 2 * padding_x, img_w - x)
        h = min(rect.h + 2 * padding_y, img_h - y)
        return Rect(x=x, y=y, w=w, h=h)

    return {"main": pad(main_rect), "vocab": pad(vocab_rect)}


def crop_image(image: Image.Image, rect: Rect) -> Image.Image:
    left = max(int(rect.x), 0)
    upper = max(int(rect.y), 0)
    right = min(int(rect.x + rect.w), image.width)
    lower = min(int(rect.y + rect.h), image.height)
    return image.crop((left, upper, right, lower))


def clean_text(text: str) -> str:
    cleaned = re.sub(r"-\n\s*", "", text)
    cleaned = re.sub(r"\s+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{2,}", "\n\n", cleaned)
    cleaned = re.sub(r"(?<!\n)\n(?!\n)", " ", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = cleaned.replace("“", '"').replace("”", '"').replace("’", "'")
    return cleaned.strip()


def split_sentences(text: str) -> list[str]:
    abbreviations = {"Mr.", "Mrs.", "Ms.", "Dr.", "Prof.", "Sr.", "Jr."}
    sentences: list[str] = []
    buffer = ""
    for token in re.split(r"(\s+)", text):
        buffer += token
        if re.search(r"[.!?][\"']?$", buffer.strip()):
            last_word = buffer.strip().split()[-1]
            if last_word in abbreviations:
                continue
            sentences.append(buffer.strip())
            buffer = ""
    if buffer.strip():
        sentences.append(buffer.strip())
    return sentences


def parse_vocab(text: str) -> list[dict[str, str]]:
    rows = []
    for line in text.splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue
        if "\t" in cleaned:
            parts = [part.strip() for part in cleaned.split("\t") if part.strip()]
        elif "/" in cleaned:
            parts = [part.strip() for part in cleaned.split("/") if part.strip()]
        elif ":" in cleaned:
            parts = [part.strip() for part in cleaned.split(":") if part.strip()]
        elif "—" in cleaned:
            parts = [part.strip() for part in cleaned.split("—") if part.strip()]
        else:
            parts = [cleaned]
        row = {
            "term": parts[0] if parts else "",
            "pos": parts[1] if len(parts) > 1 else "",
            "japanese": parts[2] if len(parts) > 2 else "",
            "extra": " ".join(parts[3:]) if len(parts) > 3 else "",
        }
        if len(parts) == 1:
            row["extra"] = ""
        rows.append(row)
    return rows


def ocr_image(image: Image.Image, lang: str) -> str:
    return pytesseract.image_to_string(image, lang=lang, config="--oem 1 --psm 6")


def job_file(job_id: str, relative_path: str) -> Path:
    return DATA_ROOT / job_id / relative_path


@app.post("/api/upload_pdf")
async def upload_pdf(file: UploadFile = File(...)) -> dict[str, Any]:
    job_id = uuid.uuid4().hex
    job_dir = ensure_job_dir(job_id)
    pdf_path = job_dir / "original.pdf"
    content = await file.read()
    pdf_path.write_bytes(content)
    pages_info = render_pdf_pages(pdf_path, job_dir)
    return {"job_id": job_id, "page_count": len(pages_info)}


@app.get("/api/jobs/{job_id}/pages")
def list_pages(job_id: str) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    page_files = sorted((job_dir / "pages").glob("page_*.png"))
    pages = []
    for index, page_path in enumerate(page_files):
        with Image.open(page_path) as img:
            pages.append(
                {
                    "page_index": index,
                    "image_url": f"/api/jobs/{job_id}/files/pages/{page_path.name}",
                    "width": img.width,
                    "height": img.height,
                }
            )
    return {"pages": pages}


@app.get("/api/jobs/{job_id}/files/{path:path}")
def serve_file(job_id: str, path: str) -> FileResponse:
    file_path = job_file(job_id, path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)


@app.post("/api/jobs/{job_id}/suggest_regions")
def suggest_regions(job_id: str, request: SuggestRegionsRequest) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    regions: dict[str, Any] = {}
    for page_index in request.page_indexes:
        page_path = job_dir / "pages" / f"page_{page_index:03d}.png"
        if not page_path.exists():
            raise HTTPException(status_code=404, detail=f"Page {page_index} not found")
        with Image.open(page_path) as img:
            suggested = suggest_regions_for_image(img)
        regions[str(page_index)] = {
            "main": suggested["main"].model_dump(),
            "vocab": suggested["vocab"].model_dump(),
        }
    return {"regions": regions}


@app.post("/api/jobs/{job_id}/ocr")
def run_ocr(job_id: str, request: OCRRequest) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    pages_output = []
    lang = request.language_hint or "eng"
    for page in request.pages:
        page_path = job_dir / "pages" / f"page_{page.page_index:03d}.png"
        if not page_path.exists():
            raise HTTPException(status_code=404, detail=f"Page {page.page_index} not found")
        with Image.open(page_path) as img:
            main_crop = crop_image(img, page.main_rect)
            vocab_crop = crop_image(img, page.vocab_rect)
            main_text_raw = ocr_image(main_crop, lang)
            vocab_text_raw = ocr_image(vocab_crop, lang)
        main_text_clean = clean_text(main_text_raw)
        sentences = split_sentences(main_text_clean)
        vocab_rows = parse_vocab(vocab_text_raw)
        pages_output.append(
            {
                "page_index": page.page_index,
                "main_text_raw": main_text_raw,
                "main_text_clean": main_text_clean,
                "sentences": sentences,
                "vocab_rows": vocab_rows,
                "warnings": [],
            }
        )

    results_path = job_dir / "ocr" / "results.json"
    results_path.parent.mkdir(parents=True, exist_ok=True)
    results_path.write_text(json.dumps({"pages": pages_output}, ensure_ascii=False, indent=2))

    return {"pages": pages_output}


@app.post("/api/jobs/{job_id}/export/word")
def export_word(job_id: str, request: ExportWordRequest) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    document = Document()
    section = document.sections[0]
    width_mm = request.layout.get("width_mm", 353) if request.layout else 353
    height_mm = request.layout.get("height_mm", 250) if request.layout else 250
    section.page_width = Mm(width_mm)
    section.page_height = Mm(height_mm)

    margins = request.layout or {}
    section.left_margin = Mm(margins.get("left_margin_mm", 15))
    section.right_margin = Mm(margins.get("right_margin_mm", 15))
    section.top_margin = Mm(margins.get("top_margin_mm", 15))
    section.bottom_margin = Mm(margins.get("bottom_margin_mm", 15))

    if request.title:
        document.add_heading(request.title, level=1)

    for page in request.pages:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        left_cell.width = Mm(220)
        right_cell.width = Mm(70)

        vocab_rows = page.get("vocab_rows", [])
        if vocab_rows:
            vocab_table = left_cell.add_table(rows=1, cols=3)
            vocab_table.style = "Table Grid"
            header = vocab_table.rows[0].cells
            header[0].text = "Term"
            header[1].text = "POS"
            header[2].text = "Japanese"
            for row in vocab_rows:
                row_cells = vocab_table.add_row().cells
                row_cells[0].text = str(row.get("term", ""))
                row_cells[1].text = str(row.get("pos", ""))
                row_cells[2].text = str(row.get("japanese", ""))

        left_cell.add_paragraph(page.get("main_text_clean", ""))
        right_cell.text = ""
        document.add_paragraph("")

    output_path = job_dir / "exports" / "handout.docx"
    document.save(output_path)
    return {"download_url": f"/api/jobs/{job_id}/files/exports/{output_path.name}"}


@app.post("/api/jobs/{job_id}/export/excel")
def export_excel(job_id: str, request: ExportExcelRequest) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Vocab"
    sheet.append(["Term", "POS", "Japanese", "Extra"])
    for row in request.vocab_rows_all:
        sheet.append([
            row.get("term", ""),
            row.get("pos", ""),
            row.get("japanese", ""),
            row.get("extra", ""),
        ])
    sheet.freeze_panes = "A2"
    output_path = job_dir / "exports" / "vocab.xlsx"
    workbook.save(output_path)
    return {"download_url": f"/api/jobs/{job_id}/files/exports/{output_path.name}"}


@app.post("/api/jobs/{job_id}/export/slides")
def export_slides(job_id: str, request: ExportSlidesRequest) -> dict[str, Any]:
    job_dir = ensure_job_dir(job_id)
    output_path = job_dir / "exports" / "slides.json"
    output_path.write_text(
        json.dumps(
            {"sentences": request.sentences_all, "style": request.style, "google": request.google},
            ensure_ascii=False,
            indent=2,
        )
    )
    return {"slide_url": f"/api/jobs/{job_id}/files/exports/{output_path.name}", "file_id": "local"}
